"""Export and reporting (Task C4).

Three export paths:
    PNG/SVG  - a single chart, via Plotly + kaleido
    PDF      - a formatted report, via reportlab
    DOCX     - the same report as Word, via python-docx

Every exported report carries the four things the brief requires: dataset
metadata, the filters that were applied, the AI-generated narrative, and the
chart image. Without the filter list a report is actively misleading - a reader
cannot tell whether "$2.3M revenue" is the whole business or one region in one
year - so the applied filters are printed even when none are set ("None").
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime

import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

BRAND = colors.HexColor("#2a78d6")
INK = colors.HexColor("#0b0b0b")
MUTED = colors.HexColor("#52514e")
RULE = colors.HexColor("#e1e0d9")


@dataclass
class ReportPayload:
    """Everything that goes into an exported report."""

    title: str
    question: str
    narrative: str
    data: pd.DataFrame
    figure: go.Figure | None = None
    code: str = ""
    caption: str = ""
    filters: dict[str, object] = field(default_factory=dict)
    dataset_name: str = "Global E-Commerce Sales (Superstore)"
    row_count: int = 0
    generated_at: datetime = field(default_factory=datetime.now)
    model: str = ""

    def filter_lines(self) -> list[str]:
        """Applied filters, rendered for print. Never silently empty."""
        if not self.filters:
            return ["None — the full dataset was used."]
        lines = []
        for key, value in self.filters.items():
            if isinstance(value, (list, tuple)):
                value = ", ".join(str(v) for v in value) if value else "all"
            lines.append(f"{key}: {value}")
        return lines


# --------------------------------------------------------------------- charts
def figure_to_png(figure: go.Figure, width: int = 1000, height: int = 560, scale: int = 2) -> bytes:
    """Render a Plotly figure to PNG bytes.

    Requires kaleido. If it is missing the caller gets a clear error rather than
    a corrupt file, because a silently chart-less report is worse than a failed
    export.
    """
    try:
        return figure.to_image(format="png", width=width, height=height, scale=scale)
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(
            "Could not render the chart to PNG. Is `kaleido` installed? "
            f"(pip install kaleido). Underlying error: {exc}"
        ) from exc


def figure_to_svg(figure: go.Figure, width: int = 1000, height: int = 560) -> bytes:
    """Render a Plotly figure to SVG bytes."""
    try:
        return figure.to_image(format="svg", width=width, height=height)
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Could not render the chart to SVG: {exc}") from exc


# ----------------------------------------------------------------- markdown
_MD_BOLD = re.compile(r"\*\*(.+?)\*\*")
_MD_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_MD_CODE = re.compile(r"`(.+?)`")


def _md_to_rl(text: str) -> str:
    """Convert the LLM's Markdown to reportlab's mini-HTML.

    Only the inline marks the narrative prompt actually produces are handled -
    bold, italic, and code. A full Markdown parser would be dead weight.
    """
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = _MD_BOLD.sub(r"<b>\1</b>", text)
    text = _MD_ITALIC.sub(r"<i>\1</i>", text)
    text = _MD_CODE.sub(r'<font face="Courier">\1</font>', text)
    return text


def _narrative_blocks(narrative: str, body, bullet) -> list:
    """Split the narrative into reportlab flowables, preserving bullet lists."""
    blocks = []
    for raw in narrative.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith(("- ", "* ", "• ")):
            blocks.append(Paragraph(_md_to_rl(line[2:].strip()), bullet))
        elif line.startswith("#"):
            heading = line.lstrip("#").strip()
            blocks.append(Paragraph(f"<b>{_md_to_rl(heading)}</b>", body))
        else:
            blocks.append(Paragraph(_md_to_rl(line), body))
    return blocks


# ---------------------------------------------------------------------- PDF
def to_pdf(payload: ReportPayload) -> bytes:
    """Render the report as a PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=payload.title,
        author="AI Analytics Platform",
    )

    sheet = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "H1", parent=sheet["Heading1"], fontSize=18, leading=23,
        textColor=INK, spaceAfter=4, alignment=TA_LEFT,
    )
    h2 = ParagraphStyle(
        "H2", parent=sheet["Heading2"], fontSize=12, leading=16,
        textColor=BRAND, spaceBefore=14, spaceAfter=6,
    )
    body = ParagraphStyle(
        "Body", parent=sheet["BodyText"], fontSize=10, leading=15, textColor=INK,
    )
    small = ParagraphStyle(
        "Small", parent=sheet["BodyText"], fontSize=8.5, leading=12, textColor=MUTED,
    )
    bullet = ParagraphStyle(
        "Bullet", parent=body, leftIndent=12, bulletIndent=2, spaceBefore=2,
        bulletText="•",
    )

    story: list = [
        Paragraph(payload.title, h1),
        Paragraph(
            f"Generated {payload.generated_at:%d %B %Y, %H:%M} · "
            f"{payload.dataset_name} · {payload.row_count:,} rows in scope"
            + (f" · model: {payload.model}" if payload.model else ""),
            small,
        ),
        Spacer(1, 10),
    ]

    story.append(Paragraph("Question", h2))
    story.append(Paragraph(_md_to_rl(payload.question), body))

    story.append(Paragraph("Applied filters", h2))
    for line in payload.filter_lines():
        story.append(Paragraph(_md_to_rl(line), bullet))

    story.append(Paragraph("AI-generated analysis", h2))
    story.extend(_narrative_blocks(payload.narrative, body, bullet))

    if payload.figure is not None:
        story.append(Paragraph("Chart", h2))
        try:
            png = figure_to_png(payload.figure, width=900, height=500)
            story.append(Image(io.BytesIO(png), width=16 * cm, height=8.9 * cm))
            if payload.caption:
                story.append(Spacer(1, 4))
                story.append(Paragraph(_md_to_rl(payload.caption), small))
        except RuntimeError as exc:
            story.append(Paragraph(f"[Chart unavailable: {exc}]", small))

    if not payload.data.empty:
        story.append(Paragraph("Result data", h2))
        preview = payload.data.head(25)
        table_data = [list(preview.columns)] + [
            [f"{v:,.2f}" if isinstance(v, float) else str(v) for v in row]
            for row in preview.itertuples(index=False)
        ]
        table = Table(table_data, repeatRows=1, hAlign="LEFT")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), BRAND),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("GRID", (0, 0), (-1, -1), 0.4, RULE),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9f9f7")]),
                    ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
                ]
            )
        )
        story.append(table)
        if len(payload.data) > 25:
            story.append(Spacer(1, 4))
            story.append(
                Paragraph(f"Showing 25 of {len(payload.data):,} rows.", small)
            )

    if payload.code:
        story.append(Paragraph("Generated query", h2))
        code_style = ParagraphStyle(
            "Code", parent=small, fontName="Courier", fontSize=8, textColor=INK
        )
        for line in payload.code.split("\n"):
            story.append(Paragraph(_md_to_rl(line), code_style))

    doc.build(story)
    return buffer.getvalue()


# --------------------------------------------------------------------- DOCX
def to_docx(payload: ReportPayload) -> bytes:
    """Render the report as a Word document with styled headings and tables."""
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    document.add_heading(payload.title, level=0)

    meta = document.add_paragraph()
    meta_run = meta.add_run(
        f"Generated {payload.generated_at:%d %B %Y, %H:%M}  |  "
        f"{payload.dataset_name}  |  {payload.row_count:,} rows in scope"
        + (f"  |  model: {payload.model}" if payload.model else "")
    )
    meta_run.font.size = Pt(8.5)
    meta_run.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_heading("Question", level=1)
    document.add_paragraph(payload.question)

    document.add_heading("Applied filters", level=1)
    for line in payload.filter_lines():
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("AI-generated analysis", level=1)
    for raw in payload.narrative.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith(("- ", "* ", "• ")):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(paragraph, line[2:].strip())
        elif line.startswith("#"):
            document.add_heading(line.lstrip("#").strip(), level=2)
        else:
            paragraph = document.add_paragraph()
            _add_markdown_runs(paragraph, line)

    if payload.figure is not None:
        document.add_heading("Chart", level=1)
        try:
            png = figure_to_png(payload.figure, width=900, height=500)
            document.add_picture(io.BytesIO(png), width=Inches(6.3))
            if payload.caption:
                caption = document.add_paragraph()
                run = caption.add_run(payload.caption)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)
        except RuntimeError as exc:
            document.add_paragraph(f"[Chart unavailable: {exc}]")

    if not payload.data.empty:
        document.add_heading("Result data", level=1)
        preview = payload.data.head(25)
        table = document.add_table(rows=1, cols=len(preview.columns))
        table.style = "Light Grid Accent 1"

        for cell, column in zip(table.rows[0].cells, preview.columns):
            cell.text = str(column)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for row in preview.itertuples(index=False):
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.text = f"{value:,.2f}" if isinstance(value, float) else str(value)

        if len(payload.data) > 25:
            note = document.add_paragraph()
            run = note.add_run(f"Showing 25 of {len(payload.data):,} rows.")
            run.font.size = Pt(8.5)
            run.italic = True

    if payload.code:
        document.add_heading("Generated query", level=1)
        code_paragraph = document.add_paragraph()
        code_run = code_paragraph.add_run(payload.code)
        code_run.font.name = "Consolas"
        code_run.font.size = Pt(9)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_markdown_runs(paragraph, text: str) -> None:
    """Add text to a Word paragraph, honouring **bold** segments."""
    for i, part in enumerate(_MD_BOLD.split(text)):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1  # split() alternates outside/inside the markers
